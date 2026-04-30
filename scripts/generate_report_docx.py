"""
scripts/generate_report_docx.py
================================
Gera um relatorio Word (.docx) descrevendo o pipeline de identificacao
e os resultados obtidos no batch (outputs/results.json).

Uso:
  python -m scripts.generate_report_docx  \
      --results outputs/results.json     \
      --output  outputs/relatorio.docx
"""
from __future__ import annotations

import argparse
import json
import statistics
from collections import Counter
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ════════════════════════════════════════════════════════════════════
# Helpers de estatistica
# ════════════════════════════════════════════════════════════════════

CONF_HIGH = 0.80
CONF_MED  = 0.55


def _confidence_band(score: float) -> str:
    if score >= CONF_HIGH:
        return "alta"
    if score >= CONF_MED:
        return "media"
    return "baixa"


def _short_source(path: str) -> str:
    """Extrai um identificador legivel da fonte (livro/artigo) a partir do path."""
    p = Path(path)
    # paths de PDF: _pdf_images/<stem>/<arquivo>.jpeg
    parts = p.parts
    if "_pdf_images" in parts:
        idx = parts.index("_pdf_images")
        if idx + 1 < len(parts):
            stem = parts[idx + 1]
            return stem[:60] + ("..." if len(stem) > 60 else "")
    return p.parent.name or p.name


def compute_stats(results: list[dict]) -> dict:
    top1_scores = [r["matches"][0]["score"] for r in results if r["matches"]]
    top5_means  = [
        sum(m["score"] for m in r["matches"]) / len(r["matches"])
        for r in results if r["matches"]
    ]

    bands = Counter(_confidence_band(s) for s in top1_scores)
    sources = Counter(_short_source(r["matches"][0]["path"])
                      for r in results if r["matches"])

    return {
        "n_samples":      len(results),
        "top1_min":       min(top1_scores),
        "top1_max":       max(top1_scores),
        "top1_mean":      statistics.mean(top1_scores),
        "top1_median":    statistics.median(top1_scores),
        "top1_stdev":     statistics.stdev(top1_scores) if len(top1_scores) > 1 else 0.0,
        "top5_mean_avg":  statistics.mean(top5_means),
        "bands":          dict(bands),
        "top_sources":    sources.most_common(5),
        "top1_scores":    top1_scores,
        "results":        results,
    }


# ════════════════════════════════════════════════════════════════════
# Graficos
# ════════════════════════════════════════════════════════════════════

def plot_score_histogram(scores: list[float], path: str):
    fig, ax = plt.subplots(figsize=(7, 3.6))
    ax.hist(scores, bins=24, color="#1D9E75", edgecolor="white")
    ax.axvline(CONF_HIGH, color="#999", linestyle="--", linewidth=1, label=f"alta (≥{CONF_HIGH})")
    ax.axvline(CONF_MED,  color="#bbb", linestyle="--", linewidth=1, label=f"media (≥{CONF_MED})")
    ax.set_xlabel("Score de similaridade (cosseno) - top-1")
    ax.set_ylabel("Numero de amostras")
    ax.set_title("Distribuicao dos scores top-1")
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(loc="upper right", fontsize=9, frameon=False)
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def plot_bands_bar(bands: dict, total: int, path: str):
    order  = ["alta", "media", "baixa"]
    counts = [bands.get(k, 0) for k in order]
    pct    = [c / total * 100 if total else 0 for c in counts]
    colors = ["#1D9E75", "#EF9F27", "#E24B4A"]

    fig, ax = plt.subplots(figsize=(7, 3.0))
    bars = ax.bar(order, counts, color=colors, width=0.55)
    for bar, c, p in zip(bars, counts, pct):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + total * 0.01,
                f"{c} ({p:.1f}%)", ha="center", va="bottom", fontsize=10)
    ax.set_ylim(0, max(counts) * 1.18 if counts else 1)
    ax.set_ylabel("Numero de amostras")
    ax.set_title(
        f"Distribuicao por faixa de confianca   "
        f"(alta ≥{CONF_HIGH}  ·  media ≥{CONF_MED}  ·  baixa <{CONF_MED})"
    )
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


# ════════════════════════════════════════════════════════════════════
# Word document
# ════════════════════════════════════════════════════════════════════

def _add_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    return h


def _add_p(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v
        for cell in (c0, c1):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in c0.paragraphs[0].runs:
            run.bold = True
    return table


def build_docx(stats: dict, output_path: str, hist_png: str, bands_png: str):
    doc = Document()

    # ── estilo padrao ─────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Titulo ────────────────────────────────────────────────────
    title = doc.add_heading("Relatorio - Identificacao de Otolitos por Similaridade Visual", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    p.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y as %H:%M')}").italic = True
    p.add_run(f"   |   Amostras analisadas: {stats['n_samples']}").italic = True

    # ── Sumario executivo ─────────────────────────────────────────
    _add_heading(doc, "1. Sumario executivo", 1)
    doc.add_paragraph(
        "Este relatorio descreve um pipeline de visao computacional para identificacao "
        "de especies de peixes a partir de fotografias de otolitos. O sistema compara "
        "cada amostra contra um banco de referencias (extraidas de atlas e artigos "
        "cientificos) e devolve as 5 referencias visualmente mais similares, com um "
        "score quantitativo. O processamento e 100% local, sem uso de APIs pagas."
    )
    doc.add_paragraph(
        f"Foram processadas {stats['n_samples']} amostras (fotografias HEIC do "
        f"acervo) contra um banco de 934 imagens de referencia extraidas de 5 "
        f"obras cientificas (Giaretta et al. 2016, Haimovici et al. 2024, Stevens "
        f"et al. 2024, MAFIS 37(1) e Otolith characterization paper)."
    )

    # ── Metodologia ──────────────────────────────────────────────
    _add_heading(doc, "2. Metodologia", 1)

    _add_heading(doc, "2.1 Visao geral do pipeline", 2)
    doc.add_paragraph(
        "Cada amostra passa pela seguinte sequencia de etapas, executadas "
        "automaticamente:"
    )
    for step in [
        "Leitura da imagem (HEIC/JPG/PNG/TIFF/WebP) - decodificacao via PIL+pillow_heif.",
        "Reducao para no maximo 1024 px no maior lado (preserva proporcao) - "
        "evita custo computacional desnecessario nas etapas seguintes; "
        "o classificador final opera em 224x224.",
        "Filtragem de ruido com kernel Gaussiano 3x3.",
        "Realce de contraste local (CLAHE - Contrast Limited Adaptive Histogram "
        "Equalization) sobre o canal L do espaco LAB - util para iluminacao "
        "irregular tipica de fotos de bancada.",
        "Segmentacao do otolito do fundo via GrabCut (OpenCV) - algoritmo "
        "iterativo de corte de grafo, sem necessidade de modelo externo.",
        "Extracao de descritor visual (embedding) com a CNN ResNet-50 "
        "pre-treinada no ImageNet, gerando um vetor de 2.048 dimensoes por imagem.",
        "Busca por similaridade no banco de referencias usando indice FAISS "
        "(IndexFlatIP) com normalizacao L2 - equivalente a similaridade do "
        "cosseno.",
        "Retorno dos 5 itens mais similares (top-5) com scores e caminhos.",
    ]:
        doc.add_paragraph(step, style="List Number")

    _add_heading(doc, "2.2 Metrica de comparacao", 2)
    doc.add_paragraph(
        "A comparacao entre uma amostra e uma referencia e feita pela "
        "similaridade do cosseno entre seus embeddings (vetores de 2.048 "
        "dimensoes extraidos pela ResNet-50). Apos normalizacao L2, ela e "
        "matematicamente identica ao produto interno entre os vetores:"
    )
    p = doc.add_paragraph()
    p.add_run("    score(A, B) = (A . B) / (||A|| * ||B||)").italic = True
    doc.add_paragraph(
        "Valores possiveis vao de -1 (vetores opostos) a +1 (identicos). "
        "Para embeddings de imagens naturais usando ResNet-50 ImageNet, "
        "scores tipicos ficam entre 0,20 e 0,90, com:"
    )
    bullets = [
        ("alta confianca", f"score ≥ {CONF_HIGH:.2f}",
         "imagem visualmente muito proxima da referencia (mesma especie ou "
         "mesma fotografia em outra pose/iluminacao)."),
        ("confianca media", f"{CONF_MED:.2f} ≤ score < {CONF_HIGH:.2f}",
         "similaridade plausivel, requer validacao por especialista."),
        ("baixa confianca", f"score < {CONF_MED:.2f}",
         "sem candidato claro - amostra possivelmente fora da cobertura "
         "do banco ou com qualidade prejudicada."),
    ]
    for nome, faixa, desc in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{nome} ").bold = True
        p.add_run(f"({faixa}): ").italic = True
        p.add_run(desc)

    _add_heading(doc, "2.3 Banco de referencias", 2)
    doc.add_paragraph(
        "O banco contem 934 imagens de otolitos extraidas automaticamente de 5 "
        "obras cientificas em PDF. As imagens foram passadas pelo mesmo pipeline "
        "de pre-processamento e tiveram seus embeddings indexados com FAISS "
        "para busca eficiente. Como as imagens vieram de figuras de PDFs "
        "(sem rotulo de especie por imagem), o campo de label retornado nos "
        "resultados aparece vazio - cada match aponta para o caminho da imagem "
        "de referencia, e o nome do artigo/livro pode ser identificado pelo "
        "diretorio de origem."
    )

    # ── Resultados ───────────────────────────────────────────────
    _add_heading(doc, "3. Resultados", 1)

    _add_heading(doc, "3.1 Estatisticas globais", 2)
    rows = [
        ("Amostras processadas",       f"{stats['n_samples']}"),
        ("Score top-1 - minimo",       f"{stats['top1_min']:.4f}"),
        ("Score top-1 - maximo",       f"{stats['top1_max']:.4f}"),
        ("Score top-1 - media",        f"{stats['top1_mean']:.4f}"),
        ("Score top-1 - mediana",      f"{stats['top1_median']:.4f}"),
        ("Score top-1 - desvio padrao", f"{stats['top1_stdev']:.4f}"),
        ("Media dos scores top-5 (por amostra, depois entre amostras)",
            f"{stats['top5_mean_avg']:.4f}"),
    ]
    _add_kv_table(doc, rows)

    _add_heading(doc, "3.2 Distribuicao por faixa de confianca", 2)
    doc.add_picture(bands_png, width=Cm(15))
    bands = stats["bands"]
    n     = stats["n_samples"]
    doc.add_paragraph(
        f"Das {n} amostras, "
        f"{bands.get('alta', 0)} ({bands.get('alta', 0)/n*100:.1f}%) cairam "
        f"em alta confianca, "
        f"{bands.get('media', 0)} ({bands.get('media', 0)/n*100:.1f}%) em "
        f"confianca media e "
        f"{bands.get('baixa', 0)} ({bands.get('baixa', 0)/n*100:.1f}%) em "
        f"baixa confianca."
    )

    _add_heading(doc, "3.3 Histograma dos scores top-1", 2)
    doc.add_picture(hist_png, width=Cm(15))
    doc.add_paragraph(
        "O histograma mostra a distribuicao dos scores top-1 entre as amostras. "
        "Concentracao em uma faixa estreita indica que o banco oferece um nivel "
        "de similaridade consistente para a maioria das amostras."
    )

    _add_heading(doc, "3.4 Fontes mais frequentes nos top-1", 2)
    doc.add_paragraph(
        "Distribuicao das obras de origem entre os melhores candidatos top-1:"
    )
    src_table = doc.add_table(rows=1, cols=3)
    hdr = src_table.rows[0].cells
    hdr[0].text = "Fonte (top-1)"
    hdr[1].text = "Amostras"
    hdr[2].text = "% do total"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
    for source, count in stats["top_sources"]:
        row = src_table.add_row().cells
        row[0].text = source
        row[1].text = str(count)
        row[2].text = f"{count / n * 100:.1f}%"

    _add_heading(doc, "3.5 Exemplos de amostras (top-1 por categoria de score)", 2)
    # Pega 1 amostra de cada faixa (alta/media/baixa) se houver
    by_band = {"alta": [], "media": [], "baixa": []}
    for r in stats["results"]:
        if not r["matches"]:
            continue
        s = r["matches"][0]["score"]
        by_band[_confidence_band(s)].append((s, r))

    for band in ("alta", "media", "baixa"):
        if not by_band[band]:
            continue
        # exemplo do meio (ordena por score e pega o de score do meio na faixa)
        by_band[band].sort()
        mid = by_band[band][len(by_band[band]) // 2][1]
        doc.add_paragraph(
            f"Faixa {band} - exemplo: {Path(mid['sample']).name}",
            style="Heading 3",
        )
        ex_table = doc.add_table(rows=1, cols=3)
        h = ex_table.rows[0].cells
        h[0].text = "Posicao"
        h[1].text = "Score"
        h[2].text = "Referencia"
        for c in h:
            c.paragraphs[0].runs[0].bold = True
        for m in mid["matches"]:
            row = ex_table.add_row().cells
            row[0].text = f"#{m['rank']}"
            row[1].text = f"{m['score']:.4f}"
            row[2].text = _short_source(m["path"]) + "  /  " + Path(m["path"]).name

    # ── Limitacoes ───────────────────────────────────────────────
    _add_heading(doc, "4. Limitacoes e consideracoes para a interpretacao", 1)
    for item in [
        "Os scores sao similaridade visual (descritor da CNN), nao taxonomia. "
        "Score alto indica imagens parecidas - confirmacao de especie deve ser "
        "feita por especialista comparando os top-K visualmente.",
        "O banco atual nao possui rotulo de especie por imagem (as referencias "
        "vieram de figuras de PDF). Para retornar nomes de especies, o usuario "
        "precisa organizar imagens em data/referencias/imagens/Nome_da_especie/ "
        "e reconstruir o banco - o pipeline ja suporta esse fluxo.",
        "A segmentacao GrabCut assume otolito centralizado e fundo razoavelmente "
        "uniforme. Em fotos com fundo complexo ou multiplos objetos, a "
        "qualidade do recorte pode comprometer o embedding. O codigo possui "
        "comentarios apontando o uso opcional de SAM (Segment Anything Model, "
        "Meta) para casos dificeis.",
        "A ResNet-50 usada e pre-treinada em ImageNet. Para ganhos de "
        "acuracia em dados taxonomicos, recomenda-se fine-tuning com imagens "
        "rotuladas (suporte ja implementado em training/finetune.py).",
        "Os 934 itens de referencia foram extraidos automaticamente de PDFs e "
        "podem incluir figuras nao-otolito (graficos, mapas, tabelas) que "
        "nao foram filtradas. Isso pode adicionar ruido aos resultados em "
        "amostras de baixa confianca.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ── Reprodutibilidade ────────────────────────────────────────
    _add_heading(doc, "5. Reprodutibilidade", 1)
    doc.add_paragraph(
        "O ambiente de execucao e definido por Dockerfile + docker-compose.yml, "
        "garantindo reprodutibilidade. As principais bibliotecas:"
    )
    rows = [
        ("Linguagem",         "Python 3.11"),
        ("Modelo CNN",        "torchvision ResNet-50 (ImageNet weights)"),
        ("Pre-processamento", "OpenCV 4.9 (denoise, CLAHE, GrabCut)"),
        ("Decodificacao HEIC", "Pillow + pillow_heif"),
        ("Indice de busca",   "FAISS-CPU 1.8 (IndexFlatIP, normalizacao L2)"),
        ("Extracao de PDFs",  "PyMuPDF 1.24"),
        ("Interface web",     "Gradio 5.x"),
        ("Aceleracao",        "CUDA opcional (suportado, GPU detectada na execucao)"),
    ]
    _add_kv_table(doc, rows)

    doc.add_paragraph(
        "Comando reproduzido para esta entrega:"
    )
    p = doc.add_paragraph()
    p.add_run(
        "    docker compose exec otolith python -m scripts.batch_identify "
        "--samples_dir data/amostras --top_k 5"
    ).font.name = "Consolas"

    # ── Anexos ───────────────────────────────────────────────────
    _add_heading(doc, "6. Anexos gerados", 1)
    for item in [
        "outputs/results.csv - 1 linha por amostra, top-5 com path/score/label.",
        "outputs/results.json - mesmos dados em JSON estruturado.",
        "outputs/report.html - relatorio visual com miniaturas das amostras e "
        "dos top-5, util para inspecao por especialista.",
        "outputs/relatorio.docx - este documento.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(output_path)


# ════════════════════════════════════════════════════════════════════
# CLI
# ════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--results", default="outputs/results.json")
    parser.add_argument("--output",  default="outputs/relatorio.docx")
    args = parser.parse_args()

    results = json.loads(Path(args.results).read_text(encoding="utf-8"))
    stats = compute_stats(results)

    out_dir = Path(args.output).parent
    out_dir.mkdir(parents=True, exist_ok=True)

    hist_png  = str(out_dir / "_hist_top1.png")
    bands_png = str(out_dir / "_bands_top1.png")

    plot_score_histogram(stats["top1_scores"], hist_png)
    plot_bands_bar(stats["bands"], stats["n_samples"], bands_png)

    build_docx(stats, args.output, hist_png, bands_png)

    print(f"\n[OK] Relatorio salvo em: {args.output}")
    print(f"     Histograma:        {hist_png}")
    print(f"     Faixas:            {bands_png}")
    print()
    print(f"Resumo: n={stats['n_samples']}  |  "
          f"top1 mean={stats['top1_mean']:.4f}  |  "
          f"alta={stats['bands'].get('alta', 0)}  "
          f"media={stats['bands'].get('media', 0)}  "
          f"baixa={stats['bands'].get('baixa', 0)}")


if __name__ == "__main__":
    main()
